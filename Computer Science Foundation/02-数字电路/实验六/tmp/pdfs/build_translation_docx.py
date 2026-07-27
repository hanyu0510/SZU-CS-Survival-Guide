from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = Path(__file__).with_name("translation_source.md")
OUTPUT_DIR = ROOT / "output" / "documents"
OUTPUT = OUTPUT_DIR / "old-exam-problems-for-mt1_题干中文翻译.docx"


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, name="Calibri", east_asia="Microsoft YaHei", size=11, color=None, bold=None):
    font = style.font
    font.name = name
    font.size = Pt(size)
    if color:
        font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[style_name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    set_style_font(doc.styles["List Bullet"], size=11)
    doc.styles["List Bullet"].paragraph_format.left_indent = Inches(0.375)
    doc.styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.188)
    doc.styles["List Bullet"].paragraph_format.space_after = Pt(4)
    doc.styles["List Bullet"].paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("题干中文翻译")
    set_run_font(run, size=9, color="666666")


def add_title(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=20, bold=True, color="0B2545")


def add_code_line(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    shade_paragraph(paragraph, "F4F6F9")
    run = paragraph.add_run(text)
    set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=9, color="333333")


def build():
    doc = Document()
    configure_document(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_code = False
    first_title = True

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            add_code_line(doc, line)
            continue
        if not line:
            continue
        if line.startswith("# "):
            if first_title:
                add_title(doc, line[2:].strip())
                first_title = False
            else:
                doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(line[2:].strip())
            set_run_font(run)
            continue
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line)
        set_run_font(run)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
