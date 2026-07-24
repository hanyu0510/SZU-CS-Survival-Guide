from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


WORKDIR = Path.cwd()
SOURCE = WORKDIR / "实验四--汉雨2024040042 (已自动恢复) (已自动恢复).docx"
OUTPUT = WORKDIR / "实验四--汉雨2024040042 (已自动恢复) (已自动恢复)-补充版.docx"


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=None, bold=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_text(paragraph, text, east_asia="宋体", latin="Times New Roman", size=None, bold=None):
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, latin=latin, size=size, bold=bold)
    return run


def insert_paragraph_before(target, text="", style=None):
    new_p = OxmlElement("w:p")
    target._p.addprevious(new_p)
    paragraph = Paragraph(new_p, target._parent)
    if style is not None:
        paragraph.style = style
    if text:
        add_text(paragraph, text)
    return paragraph


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_layout(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_borders(table, color="8A8A8A", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def format_table_text(cell, header=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        if header:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, size=9, bold=header)


def find_paragraph(paragraphs, text):
    for paragraph in paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Cannot find paragraph: {text}")


def add_algorithm_section(doc):
    process_cell = doc.tables[0].rows[2].cells[0]
    body_style = process_cell.paragraphs[2].style
    subtitle_style = find_paragraph(process_cell.paragraphs, "数学法").style
    normal_style = doc.styles["Normal"]
    target = find_paragraph(process_cell.paragraphs, "空间优化")

    items = [
        ("", normal_style),
        ("动态规划 + 二分 + 数学法", subtitle_style),
        ("在数学法中，我们已经把问题从“有 k 个鸡蛋、n 层楼最少扔几次”转换成“给定 t 次尝试、k 个鸡蛋最多能覆盖多少层楼”。如果 f(t, k) >= n，说明 t 次已经够用。", body_style),
        ("这个判断具有单调性：当尝试次数 t 增加时，可覆盖楼层数一定不会减少。因此答案可以在尝试次数 t 上做二分查找，而不是像数学法那样从 1 次、2 次、3 次逐次往上试。", body_style),
        ("判断某个 t 是否可行时，可以继续使用动态规划递推 f(t, k)=f(t-1, k-1)+f(t-1, k)+1；也可以利用它对应的组合数形式 f(t, k)=C(t, 1)+C(t, 2)+...+C(t, k)。实验中的“动态规划 + 二分 + 数学法”采用后一种判断方式，并在覆盖楼层数已经超过 n 时提前停止，避免大数继续膨胀。", body_style),
        ("以 2 个鸡蛋、100 层楼为例，普通数学法要从 t=1 开始逐次计算覆盖层数；该方法先判断中间次数是否足够，若足够就向更小次数寻找，否则向更大次数寻找，最后得到最小可行 t。因此它特别适合 10^21、10^24 甚至更大楼层数的测试。", body_style),
        ("伪代码如下：", body_style),
        ("function CanReach(t, k, n):", normal_style),
        ("    covered = 0", normal_style),
        ("    combination = 1", normal_style),
        ("    for i from 1 to min(k, t):", normal_style),
        ("        combination = combination * (t - i + 1) / i", normal_style),
        ("        covered = covered + combination", normal_style),
        ("        if covered >= n:", normal_style),
        ("            return true", normal_style),
        ("    return false", normal_style),
        ("", normal_style),
        ("function BinaryMathEggDrop(k, n):", normal_style),
        ("    low = 0, high = 1", normal_style),
        ("    while CanReach(high, k, n) == false:", normal_style),
        ("        high = high * 2", normal_style),
        ("", normal_style),
        ("    while low < high:", normal_style),
        ("        mid = (low + high) / 2", normal_style),
        ("        if CanReach(mid, k, n):", normal_style),
        ("            high = mid", normal_style),
        ("        else:", normal_style),
        ("            low = mid + 1", normal_style),
        ("    return low", normal_style),
        ("该方法中，每次可行性判断最多计算 k 个组合数项，二分查找次数为 O(log N)，所以理论时间复杂度约为 O(K log N)，空间复杂度为 O(1)；如果使用高精度整数，还需要额外考虑大整数位数带来的常数开销。", body_style),
    ]

    for text, style in items:
        paragraph = insert_paragraph_before(target, text, style)
        if style == normal_style and text:
            for run in paragraph.runs:
                set_run_font(run, east_asia="宋体", latin="Consolas", size=9)


def add_complexity_section(doc):
    result_cell = doc.tables[1].rows[0].cells[0]
    body_style = result_cell.paragraphs[3].style
    subtitle_style = result_cell.paragraphs[1].style

    title = result_cell.add_paragraph(style=subtitle_style)
    add_text(title, "理论时间复杂度比较")

    intro = result_cell.add_paragraph(style=body_style)
    add_text(intro, "为了和实验结果对应，这里用 K 表示鸡蛋数，N 表示楼层数，T 表示最终求出的最少尝试次数。各算法的理论时间复杂度对比如下：")

    rows = [
        ("算法", "理论时间复杂度", "复杂度来源", "适用特点"),
        ("蛮力法", "指数级，约 O(2^N)", "递归枚举首次投掷楼层，且不保存重复子问题，递归树会迅速膨胀。", "只适合小规模样例和原理说明。"),
        ("普通动态规划", "O(KN^2)", "共有 O(KN) 个状态，每个状态还要枚举 1 到 N 的投掷楼层。", "比蛮力法稳定，但楼层数大时增长很快。"),
        ("动态规划+二分", "O(KN log N)", "状态数仍为 O(KN)，但每个状态用二分查找代替线性枚举。", "适合中大规模，仍需要按楼层建表。"),
        ("数学法", "O(KT)", "按尝试次数递推覆盖楼层数，每增加一次尝试更新 K 个鸡蛋状态。", "当 T 远小于 N 时效率很高。"),
        ("动态规划+二分+数学法", "O(K log N)", "在尝试次数上二分，每次用组合数形式快速判断 t 次是否足够。", "最适合超大楼层数场景。"),
        ("空间优化", "O(KT)", "时间与数学法相同，只是把二维覆盖表压缩成一维数组。", "主要优势是把空间降为 O(K)。"),
    ]

    table = result_cell.add_table(rows=len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed_layout(table)
    set_table_borders(table)
    widths = [2.3, 2.7, 5.2, 3.0]

    for row_index, row_data in enumerate(rows):
        for col_index, text in enumerate(row_data):
            cell = table.rows[row_index].cells[col_index]
            cell.text = text
            set_cell_width(cell, widths[col_index])
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
            format_table_text(cell, header=(row_index == 0))

    summary = result_cell.add_paragraph(style=body_style)
    add_text(summary, "从表中可以看出，优化并不只是让代码写得更快，而是不断改变“状态怎么定义”和“转移怎么求”。蛮力法到普通 DP 解决的是重复计算问题；DP+二分解决的是单个状态枚举过慢的问题；数学法和二分数学法则进一步改变观察角度，把楼层枚举转化为覆盖能力判断，因此在极限数据规模下优势最明显。")


def update_conclusion(doc):
    conclusion_cell = doc.tables[1].rows[1].cells[0]
    for paragraph in conclusion_cell.paragraphs:
        text = paragraph.text
        old = "本次实验围绕鸡蛋掉落问题，对蛮力法、普通动态规划、动态规划加二分优化、数学递推法以及进一步优化方法进行了实现和对比。"
        new = "本次实验围绕鸡蛋掉落问题，对蛮力法、普通动态规划、动态规划加二分优化、数学递推法、动态规划+二分+数学法优化以及空间优化方法进行了实现和对比。"
        if old in text:
            paragraph.clear()
            add_text(paragraph, text.replace(old, new))
            break


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)
    add_algorithm_section(doc)
    add_complexity_section(doc)
    update_conclusion(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
